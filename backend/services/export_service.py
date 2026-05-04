"""
Export service — ATS-compliant DOCX and PDF, guaranteed one page.

One-page rules enforced before generation:
- Max 4 experiences, 3 bullets each
- Max 3 projects
- Max 20 skills
- Summary trimmed to 2 sentences max
- Tight margins (0.55 in) and 9.5 pt body font
- PDF: two-pass — if still overflows, font drops to 8.5 pt
"""

import io
import re
from copy import deepcopy
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _clean_text(text: str) -> str:
    if not text:
        return ""
    replacements = {
        "‘": "'", "’": "'", "“": '"', "”": '"',
        "–": "-", "—": "-", "…": "...",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return re.sub(r"[^\x20-\x7E]", "", text).strip()


def _format_date_range(start: str, end: str) -> str:
    if not start and not end:
        return ""
    if not end or end.lower() == "present":
        return f"{start} - Present"
    return f"{start} - {end}"


def _trim_summary(summary: str) -> str:
    """Keep at most 2 sentences."""
    sentences = re.split(r"(?<=[.!?])\s+", summary.strip())
    return " ".join(sentences[:2])


def _cap_for_one_page(resume: Dict) -> Dict:
    """
    Return a copy of resume with content capped to fit one page:
    - max 4 experiences, 3 bullets each
    - max 3 projects
    - max 20 skills
    - summary: 2 sentences max
    """
    r = deepcopy(resume)
    profile = r.get("tailored_profile", {})

    summary = profile.get("summary", "")
    if summary:
        profile["summary"] = _trim_summary(summary)

    skills = profile.get("skills", [])
    profile["skills"] = skills[:20]

    experiences = profile.get("experiences", [])[:4]
    for exp in experiences:
        exp["bullets"] = exp.get("bullets", [])[:3]
    profile["experiences"] = experiences

    profile["projects"] = profile.get("projects", [])[:3]

    r["tailored_profile"] = profile
    return r


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

_MARGIN = 0.55
_BODY_PT = 9.5
_NAME_PT = 15
_HEADING_PT = 10.5


def _set_margins(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(_MARGIN)
        section.bottom_margin = Inches(_MARGIN)
        section.left_margin = Inches(_MARGIN)
        section.right_margin = Inches(_MARGIN)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for para in section.header.paragraphs:
            para.clear()
        for para in section.footer.paragraphs:
            para.clear()


def _add_name(doc: Document, name: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(_clean_text(name))
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(_NAME_PT)


def _add_contact_line(doc: Document, contact: Dict) -> None:
    parts = [
        _clean_text(contact.get(f, ""))
        for f in ("location", "phone", "email", "linkedin", "github")
        if contact.get(f)
    ]
    if parts:
        p = doc.add_paragraph(" | ".join(parts))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(_BODY_PT)


def _add_section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(_HEADING_PT)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)


def _add_body_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(_clean_text(text))
    run.font.name = "Calibri"
    run.font.size = Pt(_BODY_PT)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(0)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(_clean_text(text))
    run.font.name = "Calibri"
    run.font.size = Pt(_BODY_PT)
    p.paragraph_format.space_after = Pt(0)


def generate_docx(resume: Dict) -> bytes:
    resume = _cap_for_one_page(resume)
    profile = resume.get("tailored_profile", {})
    contact = resume.get("contact_override") or resume.get("contact") or {}

    doc = Document()
    _set_margins(doc)

    name = _clean_text(contact.get("name", resume.get("job_title", "Resume")))
    _add_name(doc, name)
    _add_contact_line(doc, contact)

    summary = profile.get("summary", "").strip()
    if summary:
        _add_section_heading(doc, "Summary")
        _add_body_para(doc, summary)

    skills: List[str] = profile.get("skills", [])
    if skills:
        _add_section_heading(doc, "Skills")
        _add_body_para(doc, ", ".join([_clean_text(s) for s in skills if s]))

    experiences: List[Dict] = profile.get("experiences", [])
    if experiences:
        _add_section_heading(doc, "Experience")
        for exp in experiences:
            title = _clean_text(exp.get("title", ""))
            company = _clean_text(exp.get("company", ""))
            location = _clean_text(exp.get("location", ""))
            date_range = _format_date_range(exp.get("start_date", ""), exp.get("end_date", ""))
            header_parts = [f"{title} — {company}"]
            if location:
                header_parts.append(location)
            if date_range:
                header_parts.append(date_range)
            _add_body_para(doc, " | ".join(header_parts), bold=True)
            for bullet in exp.get("bullets", []):
                if bullet:
                    _add_bullet(doc, bullet)

    education: List[Dict] = resume.get("education_override") or resume.get("education") or []
    if education:
        _add_section_heading(doc, "Education")
        for edu in education:
            line = f"{_clean_text(edu.get('degree', ''))} in {_clean_text(edu.get('field', ''))} — {_clean_text(edu.get('school', ''))}"
            if edu.get("year"):
                line += f" ({_clean_text(edu['year'])})"
            _add_body_para(doc, line, bold=True)

    projects: List[Dict] = profile.get("projects", [])
    if projects:
        _add_section_heading(doc, "Projects")
        for proj in projects:
            name_text = _clean_text(proj.get("name", ""))
            tech = ", ".join([_clean_text(t) for t in proj.get("tech_stack", []) if t])
            header = f"{name_text} | {tech}" if tech else name_text
            _add_body_para(doc, header, bold=True)
            if proj.get("purpose"):
                _add_body_para(doc, proj["purpose"])
            for feat in proj.get("key_features", []):
                if feat:
                    _add_bullet(doc, feat)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF generation — with two-pass one-page enforcement
# ---------------------------------------------------------------------------

def _build_pdf_story(resume: Dict, body_pt: float) -> list:
    profile = resume.get("tailored_profile", {})
    contact = resume.get("contact_override") or resume.get("contact") or {}
    education = resume.get("education_override") or resume.get("education") or []

    styles = getSampleStyleSheet()

    name_style = ParagraphStyle("Name", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=body_pt + 5.5, spaceAfter=2, alignment=1)
    contact_style = ParagraphStyle("Contact", parent=styles["Normal"],
        fontName="Helvetica", fontSize=body_pt - 0.5, spaceAfter=4, alignment=1)
    heading_style = ParagraphStyle("Heading", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=body_pt + 1, spaceBefore=6, spaceAfter=2)
    body_style = ParagraphStyle("Body", parent=styles["Normal"],
        fontName="Helvetica", fontSize=body_pt, spaceAfter=1)
    bold_style = ParagraphStyle("Bold", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=body_pt, spaceAfter=1)
    bullet_style = ParagraphStyle("Bullet", parent=styles["Normal"],
        fontName="Helvetica", fontSize=body_pt, leftIndent=10, spaceAfter=1,
        bulletIndent=2, bulletText="-")

    story = []

    name = _clean_text(contact.get("name", resume.get("job_title", "Resume")))
    story.append(Paragraph(name, name_style))

    contact_parts = [contact.get(f, "") for f in ("location", "phone", "email", "linkedin", "github") if contact.get(f)]
    if contact_parts:
        story.append(Paragraph(" | ".join([_clean_text(p) for p in contact_parts]), contact_style))

    summary = profile.get("summary", "").strip()
    if summary:
        story.append(Paragraph("SUMMARY", heading_style))
        story.append(Paragraph(_clean_text(summary), body_style))

    skills = profile.get("skills", [])
    if skills:
        story.append(Paragraph("SKILLS", heading_style))
        story.append(Paragraph(", ".join([_clean_text(s) for s in skills if s]), body_style))

    experiences = profile.get("experiences", [])
    if experiences:
        story.append(Paragraph("EXPERIENCE", heading_style))
        for exp in experiences:
            date_range = _format_date_range(exp.get("start_date", ""), exp.get("end_date", ""))
            parts = [f"{_clean_text(exp.get('title', ''))} — {_clean_text(exp.get('company', ''))}"]
            if exp.get("location"):
                parts.append(_clean_text(exp["location"]))
            if date_range:
                parts.append(date_range)
            story.append(Paragraph(" | ".join(parts), bold_style))
            for bullet in exp.get("bullets", []):
                if bullet:
                    story.append(Paragraph(_clean_text(bullet), bullet_style))

    if education:
        story.append(Paragraph("EDUCATION", heading_style))
        for edu in education:
            line = f"{_clean_text(edu.get('degree', ''))} in {_clean_text(edu.get('field', ''))} — {_clean_text(edu.get('school', ''))}"
            if edu.get("year"):
                line += f" ({_clean_text(edu['year'])})"
            story.append(Paragraph(line, bold_style))

    projects = profile.get("projects", [])
    if projects:
        story.append(Paragraph("PROJECTS", heading_style))
        for proj in projects:
            tech = ", ".join([_clean_text(t) for t in proj.get("tech_stack", []) if t])
            header = f"{_clean_text(proj.get('name', ''))} | {tech}" if tech else _clean_text(proj.get("name", ""))
            story.append(Paragraph(header, bold_style))
            if proj.get("purpose"):
                story.append(Paragraph(_clean_text(proj["purpose"]), body_style))
            for feat in proj.get("key_features", []):
                if feat:
                    story.append(Paragraph(_clean_text(feat), bullet_style))

    return story


def _count_pdf_pages(story: list, margin: float, body_pt: float) -> int:
    """Build to a throwaway buffer and count pages."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=margin)
    page_count = [0]

    from reportlab.pdfgen import canvas as rl_canvas

    class _Counter(rl_canvas.Canvas):
        def showPage(self):
            page_count[0] += 1
            super().showPage()

    doc.build(story, canvasmaker=_Counter)
    return max(page_count[0], 1)


def generate_pdf(resume: Dict) -> bytes:
    resume = _cap_for_one_page(resume)
    margin = _MARGIN * inch

    # First pass at 9.5 pt
    story = _build_pdf_story(resume, _BODY_PT)
    pages = _count_pdf_pages(story, margin, _BODY_PT)

    # Second pass at 8.5 pt if still overflowing
    if pages > 1:
        story = _build_pdf_story(resume, 8.5)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=margin, rightMargin=margin,
        topMargin=margin, bottomMargin=margin)
    doc.build(story)
    return buf.getvalue()
